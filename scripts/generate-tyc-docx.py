"""
Genera un documento Word (.docx) con TODOS los términos y condiciones
de préstamos de JSADR para revisión del administrador.

Incluye:
1. Pagaré en Blanco (texto exacto del abogado)
2. Pagaré Diligenciado (texto legal + campos auto-llenos)
3. Carta de Instrucciones (con cláusula de firma electrónica y biométrica)
4. Cláusula Aceleratoria (compartida)
5. Datos del Acreedor
6. Sección de Firma Electrónica (declaración legal)
7. Mensajes OTP que se envían al cliente

Fuente de los textos: src/app/api/documentos/route.ts (commit b1d2b3c)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = '/home/z/my-project/download/Terminos-y-Condiciones-Prestamos-JSADR.docx'

# === Datos del acreedor (idénticos a src/app/api/documentos/route.ts) ===
DATOS_ACREEDOR = {
    'nombre': 'JOHAN SEBASTIAN ALVAREZ DEL RIO',
    'cedula': '1.214.731.649',
    'direccion': 'CALLE 92 44A 34 / ARANJUEZ',
    'ciudad': 'MEDELLÍN · ANTIOQUIA',
    'telefonos': '3103674546 - 3235949510',
    'correo': 'jsa@jsadr.com.co',
}

# === Cláusula acceleratoria (idéntica a TEXTO_CLAUSULA_ACELERATORIA) ===
CLAUSULA_ACELERATORIA = [
    "El(los) suscriptor(es) del presente pagaré reconoce(n), acepta(n) y declara(n) expresamente que el plazo concedido para el pago de la obligación se ha otorgado en beneficio exclusivo del acreedor. En consecuencia, el acreedor podrá declarar vencido anticipadamente el plazo pactado y exigir de forma inmediata la totalidad de la obligación, sin necesidad de requerimiento judicial o extrajudicial previo, constitución en mora, interpelación o notificación adicional alguna, cuando se presente cualquiera de los siguientes eventos:",

    "a) El incumplimiento en el pago de dos (2) cuotas, consecutivas o no, derivadas del presente pagaré, del plan de amortización, acuerdo de pago o cualquier documento que haga parte integral de la obligación.",

    "b) El incumplimiento de cualquiera de las obligaciones, condiciones, garantías o compromisos asumidos por el deudor con ocasión del otorgamiento del crédito.",

    "Configurado cualquiera de los eventos anteriores, el acreedor podrá declarar de pleno derecho terminado el acuerdo de pago y exigir inmediatamente la totalidad del saldo insoluto de la obligación, incluyendo, sin limitación alguna:",

    "El capital pendiente de pago. Los intereses corrientes causados y no pagados. Los intereses de mora causados. Los gastos de cobranza prejudicial. Los honorarios de abogados y gestores de cobro. Las costas y agencias en derecho que se generen con ocasión de las acciones de recuperación de la cartera. Cualquier otro concepto derivado directa o indirectamente de la presente obligación.",

    "A partir de la fecha en que se configure el incumplimiento y hasta el pago total de la obligación, las sumas adeudadas causarán intereses moratorios liquidados a la tasa máxima legal permitida y certificada para cada período por la autoridad competente en Colombia, equivalente al límite máximo autorizado por la ley y vigente al momento de hacerse exigible la obligación, o a la tasa máxima de mora que legalmente pueda cobrarse durante el tiempo que subsista el incumplimiento, sin exceder en ningún caso los límites establecidos por las normas sobre usura y demás disposiciones aplicables.",

    "El deudor acepta expresamente que la tasa de interés moratorio aplicable podrá variar en el tiempo conforme a las certificaciones expedidas por la Superintendencia Financiera de Colombia o la entidad que haga sus veces, aplicándose en cada período la tasa máxima legal vigente al momento de la liquidación de los intereses.",

    "La declaratoria de vencimiento anticipado facultará al acreedor para iniciar de manera inmediata las acciones de cobro prejudicial, ejecutivo, judicial o cualquier otro mecanismo legal tendiente a la recuperación de la totalidad de las sumas adeudadas, sin necesidad de requerimiento adicional alguno.",

    "El deudor o deudores renuncia(mos) expresamente a cualquier requerimiento previo para constituirse en mora y acepta que la sola ocurrencia de cualquiera de los eventos de incumplimiento aquí previstos hará exigible de inmediato la totalidad de la obligación, de conformidad con las disposiciones del Código Civil, el Código de Comercio y demás normas concordantes y aplicables.",
]


def set_cell_border(cell, **kwargs):
    """Bordes de celda"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '666666')
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_paragraph(doc, text, *, bold=False, italic=False, size=11, align=None,
                  color=None, space_after=6, indent=None, font='Calibri'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {
        1: RGBColor(0x1F, 0x3A, 0x5F),
        2: RGBColor(0x2C, 0x5F, 0x8F),
        3: RGBColor(0x4A, 0x6F, 0xA5),
    }
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'B0B0B0')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_field_table(doc, fields):
    """Tabla de campos auto-llenados por el sistema"""
    table = doc.add_table(rows=len(fields), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(11)
    for i, (k, v) in enumerate(fields):
        c1, c2 = table.rows[i].cells
        c1.width = Cm(5.5)
        c2.width = Cm(11)
        for cell, txt, bold in [(c1, k, True), (c2, v, False)]:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_border(cell, top=True, left=True, bottom=True, right=True)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(txt)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.bold = bold


def add_clausa_body(doc, paragraphs, indent=0.5):
    """Renderiza los párrafos de una cláusula"""
    for p in paragraphs:
        add_paragraph(doc, p, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=indent)


def main():
    doc = Document()

    # === Márgenes ===
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # === PORTADA ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run('TÉRMINOS Y CONDICIONES')
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    r.font.name = 'Calibri'

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(24)
    r = sub.add_run('PRÉSTAMOS JSADR')
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8F)
    r.font.name = 'Calibri'

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_after = Pt(80)
    r = desc.add_run('Documentos legales vigentes en el sistema\nPagaré · Carta de Instrucciones · Cláusula Aceleratoria · Firma Electrónica')
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    r.font.name = 'Calibri'

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(6)
    r = info.add_run(f'Acreedor: {DATOS_ACREEDOR["nombre"]}')
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info2.paragraph_format.space_after = Pt(6)
    r = info2.add_run(f'C.C. {DATOS_ACREEDOR["cedula"]} · {DATOS_ACREEDOR["ciudad"]}')
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    info3 = doc.add_paragraph()
    info3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info3.add_run(f'Tel: {DATOS_ACREEDOR["telefonos"]} · {DATOS_ACREEDOR["correo"]}')
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    # Salto de página después de portada
    doc.add_page_break()

    # === ÍNDICE ===
    add_heading(doc, 'ÍNDICE DE DOCUMENTOS', level=1)
    docs_index = [
        ('1.', 'Datos del Acreedor', 'Membrete que aparece en todos los documentos.'),
        ('2.', 'Pagaré en Blanco', 'Texto exacto del abogado. Campos vacíos para diligenciamiento manual.'),
        ('3.', 'Pagaré Diligenciado', 'Texto legal + datos del cliente y codeudor auto-llenados del sistema.'),
        ('4.', 'Carta de Instrucciones', 'Autorización permanente para diligenciamiento del pagaré. Incluye cláusula de firma electrónica y biométrica.'),
        ('5.', 'Cláusula Aceleratoria', 'Vencimiento anticipado, intereses de mora y costas. Compartida por pagaré y carta.'),
        ('6.', 'Sección de Firma Electrónica', 'Declaración legal de validez de firma electrónica (Ley 527 de 1999 y Decreto 1074 de 2015).'),
        ('7.', 'Mensajes OTP enviados al cliente', 'Texto del WhatsApp que recibe el cliente con el código de verificación.'),
        ('8.', 'Notas de configuración', 'Campos auto-llenados, campos manuales y dónde están definidos en el código.'),
    ]
    for num, title, desc in docs_index:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(f'{num} ')
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8F)
        r2 = p.add_run(f'{title}: ')
        r2.bold = True
        r2.font.size = Pt(11)
        r3 = p.add_run(desc)
        r3.font.size = Pt(11)
        r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_separator(doc)

    # === 1. DATOS DEL ACREEDOR ===
    add_heading(doc, '1. Datos del Acreedor', level=1)
    add_paragraph(doc,
        'Estos datos aparecen en el membrete superior de todos los documentos legales '
        '(pagaré en blanco, pagaré diligenciado, carta de instrucciones y documento combinado). '
        'Son los datos del prestamista, tomados del documento de referencia "formato solo para pagare.docx" '
        'proporcionado por el abogado. El correo se actualizó a jsa@jsadr.com.co (dominio corporativo).',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_field_table(doc, [
        ('Nombre', DATOS_ACREEDOR['nombre']),
        ('Cédula de ciudadanía', DATOS_ACREEDOR['cedula']),
        ('Dirección', DATOS_ACREEDOR['direccion']),
        ('Ciudad', DATOS_ACREEDOR['ciudad']),
        ('Teléfonos', DATOS_ACREEDOR['telefonos']),
        ('Correo electrónico', DATOS_ACREEDOR['correo']),
    ])

    add_separator(doc)

    # === 2. PAGARÉ EN BLANCO ===
    add_heading(doc, '2. Pagaré en Blanco', level=1)
    add_paragraph(doc,
        'Texto exacto del abogado (PAGARÉ 2026.docx). Deja todos los campos en blanco para '
        'diligenciamiento manual posterior por parte del acreedor. El texto legal es idéntico '
        'al pagaré diligenciado; la única diferencia es que los campos de nombres, cédulas, '
        'domicilio, fecha y valores monetarios aparecen como líneas en blanco.',
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_heading(doc, 'Texto del Pagaré en Blanco', level=2)

    add_paragraph(doc, 'PAGARÉ No. ____________________________', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    pagaré_blanco = [
        "Yo(Nosotros),",
        "______________________________________________________________ mayor(es) de edad, con domicilio en el municipio de Medellín Antioquia",
        "Domicilio: ________________________________",
        "identificado(s) como aparece(mos) al pie de mi(nuestras) firma(s), actuando en mi (nuestro) propio nombre, o en la condición indicada al píe de mi(nuestras) firma(s), declaro(amos):",
    ]
    for p in pagaré_blanco:
        add_paragraph(doc, p, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'PRIMERO: Que me(nos) obligo(amos) a pagar solidaria, indivisible, irrevocable e '
        'incondicionalmente a la orden de JOHAN SEBASTIAN ALVAREZ DEL RIO, mayor de edad, '
        'identificado con cédula de ciudadanía No. 1.214.731.649, con domicilio en la CALLE 92 44A 34, '
        'barrio Aranjuez, Medellín (Antioquia), en adelante EL ACREEDOR, o a quien represente sus derechos, '
        'el día ( _______ ) del mes de _________________ del año _________ , en sus oficinas del país o '
        'en los puntos de pago autorizados expresamente para el efecto, las siguientes sumas de dinero:',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc, 'POR CAPITAL:', bold=True, size=11)
    add_paragraph(doc, '_____________________________________________________________________________________________________________', size=11)
    add_paragraph(doc, '($ __________________________________________________) M.C.', size=11)

    add_paragraph(doc, 'POR INTERESES CAUSADOS Y NO PAGADOS:', bold=True, size=11)
    add_paragraph(doc, '($ __________________________________________________) M.C.', size=11)
    add_paragraph(doc, '___________________________________________________________________________________________________________________', size=11)

    add_paragraph(doc, 'POR OTROS CONCEPTOS:', bold=True, size=11)
    add_paragraph(doc, '($ __________________________________________________) M.C.', size=11)
    add_paragraph(doc, '_________________________________________________________________________________', size=11)

    add_paragraph(doc,
        'SEGUNDO: Que pagare(mos) intereses moratorios a la tasa máxima legalmente autorizada sobre '
        'la suma de capital insoluto.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'TERCERO: CLÁUSULA ACELERATORIA, VENCIMIENTO ANTICIPADO E INTERESES DE MORA',
        bold=True, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_clausa_body(doc, CLAUSULA_ACELERATORIA, indent=0.8)

    add_paragraph(doc,
        'CUARTO: Que acepto(amos) expresamente cualquier endoso o cesión que de este pagaré haga '
        'EL ACREEDOR reconozco(emos) desde ya al endosatario o cesionario dentro de cualquier proceso judicial.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'QUINTO: EL ACREEDOR se podrá acoger a los términos del artículo 886 del Código de Comercio para '
        'el cobro de intereses. El presente pagaré no está sujeto a la presentación para su pago, ni al aviso '
        'de rechazo, ni al protesto para todos los efectos legales y se suscribe para ser llenado por EL '
        'ACREEDOR o su representante según las instrucciones impartidas por mi(nosotros), las cuales están '
        'contenidas en la carta de autorizaciones e instrucciones adjunta al presente documento, de conformidad '
        'con lo dispuesto en el artículo 622 del Código de Comercio.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'El suscriptor declara haber suministrado voluntariamente al acreedor copia de su documento de identidad, '
        'la cual hace parte de los soportes de identificación de la presente obligación.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'Para constancia se firma en un (1) original, con destino a JOHAN SEBASTIAN ALVAREZ DEL RIO, '
        'C.C. 1.214.731.649, quien presta el dinero a los ( ___ ) días del mes de ___________ del año _____.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_separator(doc)

    # === 3. PAGARÉ DILIGENCIADO ===
    add_heading(doc, '3. Pagaré Diligenciado', level=1)
    add_paragraph(doc,
        'Texto legal idéntico al pagaré en blanco, pero con los siguientes datos auto-llenados del sistema: '
        'nombre del cliente, cédula, dirección compuesta (dirección + barrio + municipio), teléfono y correo. '
        'También se auto-llenan la fecha de suscripción (tomada de fechaDesembolso o fechaSolicitud del préstamo) '
        'y el número de pagaré (código del préstamo, ej: PREST-2026-001).',
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'IMPORTANTE: Los valores monetarios (capital, intereses causados, otros conceptos) y las tasas/saldos '
        'NO se auto-llenan. Quedan como campos en blanco (líneas) para diligenciamiento MANUAL por parte del '
        'acreedor, conforme al modelo de pagaré en blanco + carta de instrucciones del abogado.',
        size=11, italic=True, color=RGBColor(0xC0, 0x39, 0x2B),
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_heading(doc, 'Campos auto-llenados por el sistema', level=2)
    add_field_table(doc, [
        ('PAGARÉ No.', '{prestamo.codigo}  (ej: PREST-2026-001)'),
        ('Nombre del deudor', '{cliente.nombre}'),
        ('Domicilio del deudor', '{cliente.direccion} · {cliente.barrio} · {cliente.municipio}'),
        ('Día / Mes / Año', 'Tomado de prestamo.fechaDesembolso o prestamo.fechaSolicitud'),
        ('Datos del codeudor', 'Solo si prestamo.tieneCodeudor = true'),
        ('Firma electrónica', 'Imagen de la firma dibujada por el cliente + fecha/hora de firma'),
    ])

    add_heading(doc, 'Campos que quedan en blanco (diligenciamiento manual)', level=2)
    add_field_table(doc, [
        ('POR CAPITAL', 'Línea en blanco + ($ ___) M.C.'),
        ('POR INTERESES CAUSADOS Y NO PAGADOS', 'Línea en blanco + ($ ___) M.C.'),
        ('POR OTROS CONCEPTOS', 'Línea en blanco + ($ ___) M.C.'),
        ('Tasa de interés moratorio', 'Definida legalmente como "tasa máxima legalmente autorizada"'),
    ])

    add_paragraph(doc,
        'El cuerpo legal del pagaré diligenciado es IDÉNTICO al pagaré en blanco (cláusulas PRIMERO a QUINTO, '
        'incluyendo la cláusula acceleratoria completa). La única diferencia es que los campos de identidad del '
        'deudor y del codeudor aparecen diligenciados.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_separator(doc)

    # === 4. CARTA DE INSTRUCCIONES ===
    add_heading(doc, '4. Carta de Instrucciones', level=1)
    add_paragraph(doc,
        'Texto exacto del abogado (carta de instrucciones 2026.docx). Autorización permanente para que el '
        'acreedor diligencie el pagaré en blanco según las instrucciones aquí contenidas, conforme al '
        'artículo 622 del Código de Comercio. Incluye la cláusula de aceptación de firma electrónica y '
        'tratamiento de datos biométricos (cláusula 9).',
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_heading(doc, 'Texto de la Carta de Instrucciones', level=2)

    add_paragraph(doc, 'CARTA DE INSTRUCCIONES', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 'Señores', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 'AUTORIZACION E INSTRUCCIONES PERMANENTES PARA EL DILIGENCIAMIENTO DEL PAGARÉ No.',
                  bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '{prestamo.codigo}', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    carta_intro = [
        "Yo(Nosotros),",
        "{cliente.nombre} mayor(es) de edad, con domicilio en el Municipio de Medellín Antioquia Domicilio {cliente.direccion} · {cliente.barrio} · {cliente.municipio}",
        "identificado(s) como aparece(mos) al pie mi(nuestras) firma(s), actuando en mi(nuestro) propio nombre, o en la condición indicada al píe de mi(nuestras) firma(s), declaro(amos):",
    ]
    for p in carta_intro:
        add_paragraph(doc, p, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'Que de conformidad con lo dispuesto en el artículo 622 del Código de Comercio, por medio del presente '
        'documento autorizo(amos) irrevocablemente y de manera permanente a JOHAN SEBASTIAN ALVAREZ DEL RIO, '
        'C.C. 1.214.731.649, en adelante el ACREEDOR o a quien represente sus derechos, para llenar sin previo '
        'aviso los espacios en blanco y demás aspectos generales y particulares del pagaré indicado en la '
        'referencia, el cual he(mos) otorgado a su orden con espacios en blanco y del que hago(hacemos) entrega '
        'con efectos negociables, teniendo en cuenta las siguientes instrucciones:',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    clausulas_carta = [
        ("1.", "El pagaré podrá ser llenado cuando exista incumplimiento o mora en el pago de cualquier obligación a mí (nuestro) cargo, individual o conjuntamente, en los casos estipulados en la ley, en el pagaré mismo y demás documentos suscritos por mi (nosotros). Podrá también ser endosado, previo a su diligenciamiento, en razón de ser negociado cualquier derecho de crédito a mi (nuestro) cargo, individual, conjunta y solidariamente."),

        ("2.", "La fecha de vencimiento del título valor será aquella que corresponda al día en que sea llenado el pagaré. El ACREEDOR determinará la fecha de vencimiento del Pagaré y esta corresponderá a un día cierto, de tal manera que a partir de la misma serán exigibles de inmediato todas las obligaciones contenidas en el Pagaré materia de estas instrucciones."),

        ("3.", "El espacio relacionado con el valor de capital se llenará con el monto de todas las sumas que por concepto de saldo insoluto de capital deba (mos) al ACREEDOR, en forma separada, conjunta y solidaria, el día en que sean diligenciados los espacios en blanco, conforme a la liquidación que el ACREEDOR efectúe, derivadas de todas las obligaciones exigibles a mí(nuestro) cargo y a favor del ACREEDOR, en especial la correspondiente al mutuo que hemos recibido de parte del ACREEDOR."),

        ("4.", "El espacio relacionado con los intereses causados y no pagados será el que corresponda por este concepto, tanto de intereses de plazo como de mora, derivados de las obligaciones a mi(nuestro) cargo, conforme a la liquidación que el ACREEDOR efectúe."),

        ("5.", "El espacio relacionado con el valor de otros conceptos se llenará con el monto de todas las sumas que por cualquier otro concepto yo(nosotros) deba(amos) al ACREEDOR sin atención a su naturaleza o fuente, en especial las relacionadas con los siguientes rubros: (i) Los valores que por mí(nuestra) cuenta haya cancelado el ACREEDOR, por concepto de prima(s) de los seguros que se hayan contratado por mí(nuestra) cuenta. (ii) El monto de cualquier gasto pagado por el ACREEDOR por mi(nuestra) cuenta, especialmente impuestos, timbre, honorarios de abogados, comisiones, gastos administrativos y de cobranzas, así como cualquier otra suma que se deba por concepto distinto de intereses, salvo aquellos intereses que sea permitido capitalizar."),

        ("6.", "En el evento de que en desarrollo de esta facultad se cometieren errores involuntarios en el diligenciamiento del pagaré, el ACREEDOR queda expresamente facultado para aclararlos, enmendarlos y corregirlos de manera tal que el mismo responda a sus exigencias legales."),

        ("7.", "En caso de incumplimiento, retardo o existencia de cualquier causal de aceleración contemplada en los pagarés, contratos y reglamentos, frente a cualquiera de las obligaciones a mi(nuestro) cargo, el ACREEDOR queda autorizado para acelerar el vencimiento y exigir anticipadamente el valor de las demás obligaciones de las que sea (amos) deudor(es), garante(s) o avalista(s), individual, conjunta o solidariamente, sin necesidad de requerimiento judicial o extrajudicial para constituir en mora, así como para incorporarlas al Pagaré."),

        ("8.", "Así mismo, autorizo(amos) diligenciar los espacios en blanco correspondientes al número del pagaré, el cual corresponderá a aquel que le asigne el Banco y que identifique cualquiera de las obligaciones a mi(nuestro) cargo; así como al de mi(nuestro) domicilio, mi(nuestro) nombre y dirección. Declaro(amos) expresamente haber recibido copia del presente documento para todos los efectos legales."),
    ]

    for num, texto in clausulas_carta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.3
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.5)
        r1 = p.add_run(f'{num} ')
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r2 = p.add_run(texto)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'

    # === Cláusula 9 (la importante: firma electrónica y biométrica) ===
    add_paragraph(doc,
        '9. CLÁUSULA DE ACEPTACIÓN DE FIRMA ELECTRÓNICA Y TRATAMIENTO DE DATOS BIOMÉTRICOS:',
        bold=True, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5, color=RGBColor(0xC0, 0x39, 0x2B))

    add_paragraph(doc,
        'El DEUDOR (y el CODEUDOR, si aplica) acepta, manifiesta y reconoce de manera expresa y voluntaria '
        'que el presente Pagaré y su respectiva Carta de Instrucciones son suscritos mediante el mecanismo de '
        'firma electrónica de la plataforma JSADR (Johan Sebastián Álvarez Del Río). Las partes acuerdan que '
        'dicho mecanismo sustituye la firma manuscrita, otorgándole plena validez, autenticidad, integridad y '
        'fuerza ejecutiva al título valor aquí constituido, de conformidad con el artículo 7 de la Ley 527 de '
        '1999 y el Decreto 1074 de 2015.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'Asimismo, el DEUDOR autoriza de forma previa, explícita e informada al ACREEDOR para la captura, '
        'almacenamiento y tratamiento de su dato personal sensible consistente en el registro fotográfico de '
        'su rostro sosteniendo su documento de identidad (Cédula de Ciudadanía). Esta validación biométrica se '
        'realiza con la única finalidad de verificar la identidad del firmante, mitigar riesgos de suplantación '
        'y servir como prueba de autoría de la firma electrónica, garantizando en todo momento los derechos de '
        'confidencialidad y hábeas data consagrados en la Ley 1581 de 2012.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    # === Cláusula 10 ===
    add_paragraph(doc,
        '10. El deudor autoriza expresa e irrevocablemente al acreedor para exigir, custodiar y conservar copia '
        'de su documento de identidad firmada y con impresión de huella dactilar, la cual hará parte integral de '
        'los documentos soporte de la obligación. El deudor reconoce que dichos documentos podrán ser utilizados '
        'como medio probatorio en procesos de cobro prejudicial, judicial, ejecutivo o cualquier actuación '
        'encaminada a la recuperación de la cartera derivada del presente pagaré, sin perjuicio de los demás '
        'medios de prueba legalmente admisibles.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_heading(doc, 'Checklist del expediente del crédito', level=3)
    add_paragraph(doc, 'En el expediente del crédito se debe soportar con:', bold=True, size=11, indent=0.5)
    checklist = [
        '☐ Copia de cédula al 150%',
        '☐ Firma sobre la copia de la cédula',
        '☐ Huella índice derecho o selfie sosteniendo la cédula (de ser firma digital)',
        '☐ Fecha de entrega',
        '☐ Firma de recibido del asesor o responsable',
    ]
    for item in checklist:
        add_paragraph(doc, item, size=11, indent=1.0)

    add_paragraph(doc,
        'El pagaré llenado conforme a estas instrucciones, será exigible inmediatamente y prestará mérito '
        'ejecutivo sin más requisitos y requerimientos. Declaro(amos) que conozco(cemos) y acepto(amos) los '
        'Reglamentos y/o Contratos de los productos, así como que he(mos) recibida copia de esta carta de '
        'instrucciones.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'Para constancia se firma a los ( ___ ) días del mes de ___________ del año _____.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_separator(doc)

    # === 5. CLÁUSULA ACELERATORIA ===
    add_heading(doc, '5. Cláusula Aceleratoria (texto completo)', level=1)
    add_paragraph(doc,
        'Esta cláusula aparece tanto en el pagaré (en blanco y diligenciado) como en la carta de instrucciones '
        '(referenciada en la cláusula 7). Es el texto legal que define el vencimiento anticipado, los intereses '
        'de mora y las costas de cobranza. Texto idéntico al PAGARÉ 2026.docx del abogado.',
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_clausa_body(doc, CLAUSULA_ACELERATORIA, indent=0.5)

    add_separator(doc)

    # === 6. SECCIÓN DE FIRMA ELECTRÓNICA ===
    add_heading(doc, '6. Sección de Firma Electrónica', level=1)
    add_paragraph(doc,
        'Esta sección aparece en cada documento generado (pagaré, carta y documento combinado) cuando el '
        'cliente ha completado el proceso de firma electrónica. Contiene la declaración legal de validez, '
        'los datos del firmante, la imagen de la firma dibujada, la foto de validación biométrica (selfie '
        'con cédula) y los metadatos de la firma (fecha, hora, IP, agente de usuario).',
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_heading(doc, 'Declaración legal que aparece en el documento firmado', level=2)
    add_paragraph(doc,
        'Este documento ha sido firmado electrónicamente de conformidad con los artículos 7 y siguientes de '
        'la Ley 527 de 1999 (sobre mensajes de datos y firmas electrónicas) y el Decreto 1074 de 2015, '
        'otorgándole plena validez jurídica, autenticidad e integridad.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_paragraph(doc,
        'La firma electrónica aquí contenida tiene la misma fuerza probatoria y efectos jurídicos que la '
        'firma manuscrita, conforme al artículo 7 de la Ley 527 de 1999. La integridad y autenticidad de '
        'este documento pueden verificarse mediante el código de verificación y sello digital adjuntos.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_heading(doc, 'Datos que se incluyen en la sección de firma', level=2)
    add_field_table(doc, [
        ('Estado de la firma', 'PENDIENTE / EN_PROGRESO / COMPLETADA'),
        ('Fecha y hora de firma', 'Tomada del campo firma.fechaFirma (zona horaria America/Bogota)'),
        ('Código OTP verificado', 'Últimos 4 dígitos del OTP validado (ej: ****1234)'),
        ('Teléfono OTP', 'Número al que se envió el OTP (enmascarado)'),
        ('Imagen de la firma', 'PNG en base64, dibujada por el cliente en el canvas'),
        ('Foto de validación biométrica', 'Selfie sosteniendo la cédula (capturada por cámara)'),
        ('IP del firmante', 'IP desde la que se completó la firma'),
        ('Agente de usuario', 'Navegador y sistema operativo del firmante'),
    ])

    add_heading(doc, 'Código de verificación y sello digital', level=2)
    add_paragraph(doc,
        'Cada documento generado incluye un código de verificación único (ej: JSADR-PAG-2026-001-AB12CD34) '
        'y un sello digital SHA-256 que permite verificar la autenticidad del documento en cualquier momento. '
        'Adicionalmente se incluye un código QR que enlaza con la URL pública de verificación.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_separator(doc)

    # === 7. MENSAJES OTP ===
    add_heading(doc, '7. Mensajes OTP enviados al cliente', level=1)
    add_paragraph(doc,
        'Estos son los mensajes de WhatsApp que recibe el cliente cuando debe verificar su identidad para '
        'firmar electrónicamente. El código OTP es de 6 dígitos, aleatorio, y expira en 5 minutos. '
        'Se envía por WhatsApp Cloud API de Meta (configurable desde Configuración Global → Integraciones).',
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    add_heading(doc, 'Mensaje OTP — Aceptación de T&C del préstamo', level=2)
    add_paragraph(doc,
        'Se envía cuando el gestor solicita la aceptación de términos y condiciones de un préstamo nuevo. '
        'Plantilla de texto libre (WhatsApp Cloud API).',
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    otp_aceptacion = """🔐 *CÓDIGO DE VERIFICACIÓN - ACEPTACIÓN DE PRÉSTAMO*

Hola *{cliente.nombre}*,

Para confirmar la aceptación de los Términos y Condiciones de tu préstamo *{prestamo.codigo}*, ingresa el siguiente código:

  >>  {OTP}  <<

⏰ El código expira en 5 minutos.
⚠️ No compartas este código con nadie."""

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(otp_aceptacion)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    # Fondo gris
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F4F4F4')
    p._p.get_or_add_pPr().append(shading)

    add_heading(doc, 'Mensaje OTP — Firma de pagaré o carta de instrucciones', level=2)
    add_paragraph(doc,
        'Se envía cuando el cliente (deudor o codeudor) debe firmar electrónicamente el pagaré y la carta '
        'de instrucciones. Plantilla de texto libre (WhatsApp Cloud API).',
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)

    otp_firma = """🔐 *CÓDIGO DE VERIFICACIÓN - FIRMA ELECTRÓNICA*

Hola *{cliente.nombre}*,

Como *{ROL}* del préstamo, necesitas firmar electrónicamente los Términos y Condiciones.

Ingresa el siguiente código para completar tu firma:

  >>  {OTP}  <<

⏰ El código expira en 5 minutos.
⚠️ No compartas este código con nadie."""

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(otp_firma)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F4F4F4')
    p._p.get_or_add_pPr().append(shading)

    add_separator(doc)

    # === 8. NOTAS DE CONFIGURACIÓN ===
    add_heading(doc, '8. Notas de configuración', level=1)

    add_heading(doc, 'Dónde están definidos los textos en el código', level=2)
    add_field_table(doc, [
        ('Pagaré en blanco', 'src/app/api/documentos/route.ts → generarPagareBlancoHTML()'),
        ('Pagaré diligenciado', 'src/app/api/documentos/route.ts → generarPagareDiligenciadoHTML()'),
        ('Carta de instrucciones', 'src/app/api/documentos/route.ts → generarCartaInstruccionesHTML()'),
        ('Documento combinado', 'src/app/api/documentos/route.ts → generarDocumentoCombinadoHTML()'),
        ('Cláusula acceleratoria', 'src/app/api/documentos/route.ts → TEXTO_CLAUSULA_ACELERATORIA'),
        ('Datos del acreedor', 'src/app/api/documentos/route.ts → DATOS_ACREEDOR'),
        ('Mensaje OTP aceptación', 'src/app/api/prestamos/[id]/aceptar-tyc-otp/route.ts (línea 509)'),
        ('Mensaje OTP firma', 'src/app/api/firma/route.ts (línea 168)'),
        ('Sección firma electrónica', 'src/app/api/documentos/route.ts → generarSeccionFirmaElectronica()'),
    ])

    add_heading(doc, 'Campos que se auto-llenan del sistema', level=2)
    auto_campos = [
        ('{prestamo.codigo}', 'Código único del préstamo, ej: PREST-2026-001'),
        ('{cliente.nombre}', 'Nombre completo del deudor'),
        ('{cliente.cedula}', 'Cédula de ciudadanía del deudor'),
        ('{cliente.direccion}', 'Dirección + barrio + municipio del deudor'),
        ('{cliente.telefono}', 'Teléfono del deudor (para envío de OTP)'),
        ('{cliente.email}', 'Correo del deudor'),
        ('{prestamo.tieneCodeudor}', 'Booleano: true si hay codeudor registrado'),
        ('{prestamo.codeudorNombre}', 'Nombre completo del codeudor (si aplica)'),
        ('{prestamo.codeudorCedula}', 'Cédula del codeudor (si aplica)'),
        ('{prestamo.fechaDesembolso}', 'Fecha usada en el pagaré (día / mes / año)'),
        ('{firma.imagenFirma}', 'Imagen base64 de la firma dibujada por el cliente'),
        ('{firma.fechaFirma}', 'Fecha y hora exactas de la firma (America/Bogota)'),
    ]
    add_field_table(doc, auto_campos)

    add_heading(doc, 'Documentos de referencia originales', level=2)
    add_paragraph(doc,
        'Los textos legales de este documento son copia fiel de los siguientes archivos .docx proporcionados '
        'por el abogado de JSADR y cargados al sistema como referencia:',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5)
    add_field_table(doc, [
        ('PAGARÉ 2026.docx', 'Plantilla del pagaré en blanco con el texto legal completo'),
        ('carta de instrucciones 2026.docx', 'Plantilla de la carta de instrucciones'),
        ('formato solo para pagare.docx', 'Encabezado/membrete con datos del acreedor'),
    ])

    add_paragraph(doc,
        'Estos archivos se encuentran referenciados en el código fuente del sistema (src/app/api/documentos/route.ts) '
        'y el texto legal de los documentos generados por JSADR es idéntico al de las plantillas del abogado. '
        'Cualquier modificación a estos textos debe hacerse en el código fuente y requiere verificación legal '
        'antes de ponerse en producción.',
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.5, italic=True,
        color=RGBColor(0x60, 0x60, 0x60))

    # Guardar
    doc.save(OUTPUT_PATH)
    print(f'✅ Documento generado: {OUTPUT_PATH}')
    import os
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f'   Tamaño: {size_kb:.1f} KB')


if __name__ == '__main__':
    main()
